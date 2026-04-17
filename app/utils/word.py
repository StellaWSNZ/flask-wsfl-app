import os
import re
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from num2words import num2words

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from app.report_utils.CHT_CircleProportions import circle_plot_png
from app.report_utils.helpers import load_ppmori_fonts
from app.utils.database import get_db_engine


WSFL_BLUE = "1A427D"
WSFL_LIGHT_BLUE = "EDF3FB"


# =============================================================================
# Word helpers
# =============================================================================
def close_word(DEBUG=False):
    if DEBUG:
        subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"])
    else:
        subprocess.run(
            ["taskkill", "/f", "/im", "WINWORD.EXE"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text,
    bold=False,
    font_name="PP Mori",
    font_size=9,
    font_color="000000",
    align="left"
):
    cell.text = ""
    p = cell.paragraphs[0]

    if align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor.from_string(font_color)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_title(document, title, subtitle=None):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    r = p.add_run(title)
    r.bold = True
    r.font.name = "PP Mori"
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(WSFL_BLUE)

    if subtitle:
        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_after = Pt(8)

        r2 = p2.add_run(subtitle)
        r2.font.name = "PP Mori"
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor.from_string("666666")


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)

    r = p.add_run(text)
    r.bold = True
    r.font.name = "PP Mori"
    r.font.size = Pt(14 if level == 1 else 11)
    r.font.color.rgb = RGBColor.from_string(WSFL_BLUE)
    return document


def add_body_paragraph(document, text, space_after=6):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    r = p.add_run(text)
    r.font.name = "PP Mori"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("000000")
    return document


def set_up_doc(title, subtitle=None):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_title(document, title, subtitle=subtitle)
    return document


def fmt_num(n):
    if n < 10:
        return num2words(n)
    return str(f"{n:,}")


def format_value(val, col_name=None):
    if pd.isna(val):
        return ""

    if col_name and "Percentage" in col_name:
        return f"{float(val):.1f}%"

    if isinstance(val, float):
        if float(val).is_integer():
            return f"{int(val):,}"
        return f"{val:,.1f}"

    if isinstance(val, int):
        return f"{val:,}"

    return str(val)


def prettify_columns(df):
    df = df.copy()
    df.columns = [
        re.sub(r"(?<=[a-z])(?=[A-Z])", " ", col).replace("Eqi", "EQI")
        for col in df.columns
    ]
    return df


def draw_table_df(document, df, shading=False, DEBUG=False):
    df = prettify_columns(df)

    nrow, ncol = df.shape
    table = document.add_table(rows=nrow + 1, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    if DEBUG:
        print(f"Headers: {list(df.columns)}")

    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        set_cell_shading(cell, WSFL_BLUE)
        set_cell_text(
            cell,
            col,
            bold=True,
            font_name="PP Mori",
            font_size=9,
            font_color="FFFFFF",
            align="center"
        )

    for i in range(nrow):
        for j in range(ncol):
            val = df.iat[i, j]
            col_name = df.columns[j]
            text = format_value(val, col_name)

            cell = table.cell(i + 1, j)

            if shading and i % 2 == 1:
                set_cell_shading(cell, WSFL_LIGHT_BLUE)

            align = "left"
            if "Percentage" in col_name:
                align = "right"
            elif isinstance(val, (int, float)) and not pd.isna(val):
                align = "right"

            set_cell_text(
                cell,
                text,
                bold=False,
                font_name="PP Mori",
                font_size=9,
                font_color="000000",
                align=align
            )

    return document


def build_region_phrase(region_names):
    region_names = (region_names or "").replace(" Region", "").strip()

    if not region_names:
        return ""

    parts = [x.strip() for x in region_names.split(",") if x.strip()]

    if len(parts) == 1:
        return parts[0]
    if len(parts) == 2:
        return f"{parts[0]} and {parts[1]}"
    return f"{', '.join(parts[:-1])} and {parts[-1]}"


# =============================================================================
# Chart helpers
# =============================================================================
def render_assessment_circle_chart(bucket, out_path="assessment_circle_chart.png"):
    font_family = load_ppmori_fonts("app/static/fonts")

    fig, ax = plt.subplots(figsize=(6.8, 2.7), dpi=300)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    fig.subplots_adjust(left=0.03, right=0.97, top=0.92, bottom=0.08)

    circle_plot_png(
        ax,
        stats=bucket,
        fontfamily=font_family,
        top_y=0.92,
        max_radius=None,
        gap_between=0.03,
        side_margin=0.02,
        label_gap=0.02,
        title_reserved=0.02,
        show_pct=False,
        label_box_height=0.08,
        label_fontsize=10.5,
        bottom_margin=0.03,
        label_box_width=0.22,
    )

    fig.savefig(
        out_path,
        dpi=300,
        bbox_inches="tight",
        pad_inches=0.02,
        facecolor="white"
    )
    plt.close(fig)
    return out_path


def add_assessment_circle_chart(document, bucket, DEBUG=False):
    img_path = render_assessment_circle_chart(bucket)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.2))

    if DEBUG:
        print(f"Inserted circle chart image: {img_path}")

    return document


def render_level_score_chart(df_level_total, out_path="assessment_score_chart.png"):
    font_family = load_ppmori_fonts("app/static/fonts")

    df_plot = df_level_total.pivot(
        index="TotalScore",
        columns="LevelName",
        values="AssessmentCount"
    ).fillna(0)

    level_order = [
        "Foundational Level",
        "Intermediate Level",
        "Expert Level"
    ]
    df_plot = df_plot.reindex(columns=level_order, fill_value=0)

    colours = {
        "Foundational Level": "#BBE6E9",
        "Intermediate Level": "#2EBDC2",
        "Expert Level": "#1A427D"
    }

    fig, ax = plt.subplots(figsize=(6.8, 2.5), dpi=300)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    bottom = None
    for level in level_order:
        vals = df_plot[level]
        ax.bar(
            df_plot.index,
            vals,
            bottom=bottom,
            label=level,
            color=colours[level],
            edgecolor="none",
            width=0.8
        )
        bottom = vals if bottom is None else bottom + vals

    ax.set_xlabel("Total Assessment Score", fontname=font_family, fontsize=10)
    ax.set_ylabel("Number of Assessments", fontname=font_family, fontsize=10)

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", alpha=0.2)

    ax.tick_params(axis="both", labelsize=9)

    legend = ax.legend(
        frameon=False,
        fontsize=8,
        ncol=3,
        loc="upper center",
        bbox_to_anchor=(0.5, 1.18)
    )
    for txt in legend.get_texts():
        txt.set_fontname(font_family)

    fig.tight_layout()

    fig.savefig(
        out_path,
        dpi=300,
        bbox_inches="tight",
        pad_inches=0.02,
        facecolor="white"
    )
    plt.close(fig)

    return out_path


def add_level_score_chart(document, df_level_total, DEBUG=False):
    img_path = render_level_score_chart(df_level_total)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.4))

    if DEBUG:
        print(f"Inserted score chart image: {img_path}")

    return document


def render_funder_level_chart(df_funder_levels, out_path="funder_level_chart.png"):
    font_family = load_ppmori_fonts("app/static/fonts")

    df_plot = df_funder_levels.pivot(
        index="Funder",
        columns="LevelName",
        values="AssessmentCount"
    ).fillna(0)

    level_order = [
        "Foundational Level",
        "Intermediate Level",
        "Expert Level"
    ]
    df_plot = df_plot.reindex(columns=level_order, fill_value=0)

    colours = {
        "Foundational Level": "#BBE6E9",
        "Intermediate Level": "#2EBDC2",
        "Expert Level": "#1A427D"
    }

    fig, ax = plt.subplots(figsize=(6.8, 3.2), dpi=300)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    bottom = None
    x = range(len(df_plot.index))

    for level in level_order:
        vals = df_plot[level].values
        ax.bar(
            x,
            vals,
            bottom=bottom,
            label=level,
            color=colours[level],
            edgecolor="none",
            width=0.7
        )
        bottom = vals if bottom is None else bottom + vals

    ax.set_xticks(list(x))
    ax.set_xticklabels(
        df_plot.index,
        rotation=30,
        ha="right",
        fontsize=8,
        fontname=font_family
    )
    ax.set_ylabel("Number of Assessments", fontname=font_family, fontsize=10)

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", alpha=0.2)

    ax.tick_params(axis="y", labelsize=9)

    legend = ax.legend(
        frameon=False,
        fontsize=8,
        ncol=3,
        loc="upper center",
        bbox_to_anchor=(0.5, 1.18)
    )
    for txt in legend.get_texts():
        txt.set_fontname(font_family)

    fig.tight_layout()

    fig.savefig(
        out_path,
        dpi=300,
        bbox_inches="tight",
        pad_inches=0.02,
        facecolor="white"
    )
    plt.close(fig)

    return out_path


def add_funder_level_chart(document, df_funder_levels, DEBUG=False):
    img_path = render_funder_level_chart(df_funder_levels)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.8))

    if DEBUG:
        print(f"Inserted funder level chart image: {img_path}")

    return document


# =============================================================================
# Data
# =============================================================================
def get_report_data(DEBUG=False):
    con = get_db_engine()

    df_summary = pd.read_sql(
        "EXEC KaiakoAnalysis @Request = 'OverallSummary'",
        con
    )
    df_level_total = pd.read_sql(
        "EXEC KaiakoAnalysis @Request = 'LevelTotalSchool'",
        con
    )
    df_funder_levels = pd.read_sql(
        "EXEC KaiakoAnalysis @Request = 'FunderLevelDistribution'",
        con
    )
    df_kaiako_targets = pd.read_sql(
        "EXEC KaiakoAnalysis @Request = 'KaiakoTargets'",
        con
    )

    if DEBUG:
        print(df_summary)
        print(df_level_total)
        print(df_funder_levels)
        print(df_kaiako_targets)

    return df_summary, df_level_total, df_funder_levels, df_kaiako_targets


def prepare_targets_table(df_kaiako_targets):
    df = df_kaiako_targets.copy()

    df = df.rename(columns={
        "KaiakoTarget": "Target",
        "KaiakoCount": "Assessments",
        "KaiakoClassCount": "KaiakoClasses",
        "AssessmentCompletePercentage": "ClassesAssessedPercentage",
        "TargetPercentage": "TargetPercentage"
    })

    df = df[
        [
            "Funder",
            "Target",
            "Assessments",
            "KaiakoClasses",
            "ClassesAssessedPercentage",
            "TargetPercentage"
        ]
    ]

    return df


def build_funder_chart_insights(df_funder_levels):
    if df_funder_levels.empty:
        return []

    df_plot = df_funder_levels.pivot(
        index="Funder",
        columns="LevelName",
        values="AssessmentCount"
    ).fillna(0)

    for col in ["Foundational Level", "Intermediate Level", "Expert Level"]:
        if col not in df_plot.columns:
            df_plot[col] = 0

    df_plot["Total"] = (
        df_plot["Foundational Level"] +
        df_plot["Intermediate Level"] +
        df_plot["Expert Level"]
    )

    nonzero = df_plot[df_plot["Total"] > 0].copy()
    if nonzero.empty:
        return []

    nonzero["ExpertPct"] = nonzero["Expert Level"] / nonzero["Total"] * 100
    nonzero["Range"] = (
        nonzero[["Foundational Level", "Intermediate Level", "Expert Level"]]
        .gt(0).sum(axis=1)
    )

    highest_expert_funder = nonzero["ExpertPct"].idxmax()
    highest_expert_pct = nonzero.loc[highest_expert_funder, "ExpertPct"]

    widest_range_funder = nonzero["Range"].idxmax()

    insights = [
        f"{highest_expert_funder} has the highest proportion of Expert-level assessments ({highest_expert_pct:.1f}%)."
    ]

    if nonzero["Range"].max() > 1:
        insights.append(
            f"{widest_range_funder} shows the broadest spread across assessment levels."
        )

    intermediate_dominant = nonzero[
        (nonzero["Intermediate Level"] >= nonzero["Foundational Level"]) &
        (nonzero["Intermediate Level"] >= nonzero["Expert Level"])
    ]

    if len(intermediate_dominant) >= 1:
        insights.append(
            "Intermediate-level assessments remain the most common pattern across several funders."
        )

    return insights[:3]


def build_target_insights(df_targets_table):
    if df_targets_table.empty:
        return []

    insights = []

    no_target = df_targets_table[
        df_targets_table["Target"].isna() | (df_targets_table["Target"] == 0)
    ]
    if not no_target.empty:
        if len(no_target) == 1:
            insights.append(
                f"{no_target.iloc[0]['Funder']} did not have a funded Kaiako target for the period."
            )
        else:
            insights.append(
                "Some organisations did not have a funded Kaiako target for the period."
            )

    low_completion = df_targets_table[
        df_targets_table["ClassesAssessedPercentage"].notna() &
        (df_targets_table["ClassesAssessedPercentage"] < 80)
    ]
    if not low_completion.empty:
        worst_completion = low_completion.sort_values("ClassesAssessedPercentage").iloc[0]
        insights.append(
            f"{worst_completion['Funder']} shows lower assessment completion across Kaiako classes "
            f"({worst_completion['ClassesAssessedPercentage']:.1f}%), suggesting that not all delivered "
            f"classes have recorded teacher assessments."
        )

    high_target = df_targets_table[
        df_targets_table["TargetPercentage"].notna() &
        (df_targets_table["TargetPercentage"] >= 100)
    ]
    if not high_target.empty:
        insights.append(
            "Some funders have met or exceeded their funded Kaiako targets."
        )

    return insights[:3]


# =============================================================================
# Page builders
# =============================================================================
def build_page1_content(document, df_summary, df_level_total, DEBUG=False):
    if df_summary.empty:
        add_heading(document, "Summary", level=1)
        add_body_paragraph(document, "No summary data was returned.")
        return document

    row = df_summary.loc[0]

    foundational_raw = int(row["FoundationalLevelCount"])
    intermediate_raw = int(row["IntermediateLevelCount"])
    expert_raw = int(row["ExpertLevelCount"])
    total_level_count = foundational_raw + intermediate_raw + expert_raw

    bucket = {
        "total": total_level_count,
        "buckets": {
            "Foundational Level": {
                "count": foundational_raw,
                "colour": "#BBE6E9"
            },
            "Intermediate Level": {
                "count": intermediate_raw,
                "colour": "#2EBDC2"
            },
            "Expert Level": {
                "count": expert_raw,
                "colour": "#1A427D"
            },
        },
    }

    funder_count_raw = int(row["FunderCount"])
    school_count_raw = int(row["SchoolCount"])
    region_count_raw = int(row["RegionCount"])
    school_count_eqi_raw = int(row["SchoolCountEQI446Plus"])
    kaiako_class_count_raw = int(row["KaiakoClassCount"])
    kaiako_assessment_count_raw = int(row["KaiakoAssessmentCount"])

    funder_count = fmt_num(funder_count_raw)
    school_count = fmt_num(school_count_raw)
    region_count = fmt_num(region_count_raw)
    region_names = build_region_phrase(row["RegionNames"])

    if school_count_raw > 0:
        target_school_pct = round((school_count_eqi_raw / school_count_raw) * 100, 1)
    else:
        target_school_pct = 0

    if kaiako_class_count_raw > 0:
        assessment_completion_pct = round(
            (kaiako_assessment_count_raw / kaiako_class_count_raw) * 100, 1
        )
    else:
        assessment_completion_pct = 0

    add_heading(document, "Summary", level=1)

    summary_text = (
        f"In this funding year we have had {funder_count} funded organisations "
        f"delivering Kaiako Led programmes. This delivery has taken place at "
        f"{school_count} different schools, with {target_school_pct}% being "
        f"in our target EQI range. We have delivered in {region_count} regions: "
        f"{region_names}. Across all Kaiako Led classes, {assessment_completion_pct}% "
        f"of classes have a teacher assessment recorded."
    )
    add_body_paragraph(document, summary_text)

    if total_level_count > 0:
        foundational_pct = round((foundational_raw / total_level_count) * 100, 1)
        intermediate_pct = round((intermediate_raw / total_level_count) * 100, 1)
        expert_pct = round((expert_raw / total_level_count) * 100, 1)

        level_map = {
            "Foundational": foundational_pct,
            "Intermediate": intermediate_pct,
            "Expert": expert_pct,
        }
        top_level = max(level_map, key=level_map.get)
        top_pct = level_map[top_level]

        level_text = (
            f"Overall, teacher assessments were most commonly recorded at the "
            f"{top_level} level ({top_pct}%), with the remaining assessments "
            f"split across the other levels."
        )
        add_body_paragraph(document, level_text)

    add_heading(document, "Assessment Level Distribution", level=2)
    add_body_paragraph(
        document,
        "The distribution of assessment levels and scores is shown below."
    )

    document = add_assessment_circle_chart(document, bucket, DEBUG=DEBUG)

    if not df_level_total.empty:
        document = add_level_score_chart(document, df_level_total, DEBUG=DEBUG)

        peaks_text = (
            "The peaks at scores of 4, 8, and 12 are expected, as these totals "
            "occur when a teacher is rated consistently across all four assessment "
            "areas. For example, a score of 8 reflects a teacher being rated "
            "satisfactory in each area, while 12 reflects consistently proficient "
            "ratings across the assessment criteria."
        )
        add_body_paragraph(document, peaks_text)

        insight_text = (
            "While these consistent ratings are expected, the spread across total "
            "scores shows variation within each level, particularly within the "
            "Intermediate group. This indicates that teachers classified at the "
            "same level may still differ in how consistently they meet the "
            "underlying assessment criteria."
        )
        add_body_paragraph(document, insight_text)

    return document


def build_page2_content(document, df_funder_levels, df_kaiako_targets, DEBUG=False):
    document.add_page_break()

    add_heading(document, "Funder Comparison", level=1)
    add_body_paragraph(
        document,
        "The chart below compares the distribution of Foundational, Intermediate, "
        "and Expert teacher assessments across funded organisations."
    )

    if not df_funder_levels.empty:
        document = add_funder_level_chart(document, df_funder_levels, DEBUG=DEBUG)

        chart_insights = build_funder_chart_insights(df_funder_levels)
        for sentence in chart_insights:
            add_body_paragraph(document, sentence)

    add_heading(document, "Delivery Against Targets", level=2)
    add_body_paragraph(
        document,
        "The table below shows delivery against funded targets for each organisation."
    )

    if not df_kaiako_targets.empty:
        df_targets_table = prepare_targets_table(df_kaiako_targets)
        document = draw_table_df(document, df_targets_table, shading=True, DEBUG=DEBUG)

        target_insights = build_target_insights(df_targets_table)
        for sentence in target_insights:
            add_body_paragraph(document, sentence)

    return document


# =============================================================================
# Main
# =============================================================================
if __name__ == "__main__":
    DEBUG = True

    title = "Teacher Assessment Analysis"
    subtitle = "Kaiako Led Programmes – Overview and Funder Comparison"
    filename = f"{title}.docx"

    close_word(DEBUG)

    df_summary, df_level_total, df_funder_levels, df_kaiako_targets = get_report_data(DEBUG=DEBUG)

    document = set_up_doc(title, subtitle=subtitle)
    document = build_page1_content(document, df_summary, df_level_total, DEBUG=DEBUG)
    document = build_page2_content(
        document,
        df_funder_levels,
        df_kaiako_targets,
        DEBUG=DEBUG
    )

    document.save(filename)
    os.startfile(filename)